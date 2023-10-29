from app import connect, app
from flask import make_response, send_from_directory
from models.minute_model import Minutes
from bson import ObjectId
from pydantic import ValidationError
from docx import Document
from bson.json_util import dumps
from bson.objectid import ObjectId
import openai
import json
from os import environ
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import urllib.parse

class minutes_controller:
    def transcribe_audio(self, audio_file_path):
        with open(audio_file_path, "rb") as audio_file:
            transcription = openai.Audio.transcribe("whisper-1", audio_file)
        return transcription["text"]

    def abstract_summary_extraction(self, transcription):
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points. Provide an apt title for the text. Give the response in the following format: Title: <title comes here> , Summary: <summary comes here>",
                },
                {"role": "user", "content": transcription},
            ],
        )
        return response["choices"][0]["message"]["content"]

    def key_points_extraction(self, transcription):
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about.",
                },
                {"role": "user", "content": transcription},
            ],
        )
        return response["choices"][0]["message"]["content"]

    def action_item_extraction(self, transcription):
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely.",
                },
                {"role": "user", "content": transcription},
            ],
        )
        return response["choices"][0]["message"]["content"]

    def sentiment_analysis(self, transcription):
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible.",
                },
                {"role": "user", "content": transcription},
            ],
        )
        return response["choices"][0]["message"]["content"]

    def meeting_minutes(self, transcription):
        abstract_summary = self.abstract_summary_extraction(transcription)
        temp = abstract_summary.split("Summary: ")
        abstract_summary = temp[1]
        title = temp[0].split("Title: ")[1]
        key_points = self.key_points_extraction(transcription)
        action_items = self.action_item_extraction(transcription)
        # sentiment = self.sentiment_analysis(transcription)
        print("TITLE: ", title)
        return {
            "title": title,
            "date": datetime.now().strftime("%m/%d/%Y, %H:%M:%S"),
            "abstract_summary": abstract_summary,
            "key_points": key_points,
            "action_items": action_items
            # "sentiment": sentiment,
        }

    def save_as_docx(self, minutes):
        doc = Document()
        # Create a paragraph for the title
        title = doc.add_paragraph("BuzzMinutes")
        # Set the font size and color for the title
        title.runs[0].font.size = Pt(24)  # You can adjust the size as needed
        title.runs[0].font.color.rgb = RGBColor(0x0C, 0x35, 0x6A)  # Blue color
        # Center align the title
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title = doc.add_paragraph(minutes["title"])
        # Set the font size and color for the title
        title.runs[0].font.size = Pt(20)  # You can adjust the size as needed
        title.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # RED color
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # add abstract summary
        doc.add_heading("Abstract Summary", level=1)
        doc.add_paragraph(minutes["abstract_summary"])
        # doc.add_paragraph()
        # add key points
        doc.add_heading("Key Points", level=1)
        doc.add_paragraph(minutes["key_points"])
        # doc.add_paragraph()
        # add action items
        doc.add_heading("Action Items", level=1)
        doc.add_paragraph(minutes["action_items"])
        # doc.add_paragraph()
        # # add key points
        # doc.add_heading("Sentiment Analysis", level=1)
        # doc.add_paragraph(minutes["sentiment_analysis"])
        # doc.add_paragraph()
        # save the document
        path = "docs/" + minutes["title"].split(":")[0]+".docx"
        doc.save(path)
        # return json
        return make_response(
            {
                "success": "true",
                "message": "file downloaded successfully",
                "data": environ.get("BASE_URL")
                + "/" + path,
            }
        )

    def getall(self):
        minutes = connect.minutes
        try:
            minutes = minutes.find()
            return make_response(
                {"success": "true","message":"all minutes retrieved successfully", "data": json.loads(dumps(list(minutes)))}, 200
            )
        except:
            return make_response({"success": "true", "message": "server error"}, 400)

    def getminutesbyid(self, id):
        try:
            minutes = connect.minutes.find({"_id": ObjectId(id)})
            return make_response(
                {
                    "success": "true",
                    "message": "minute retrieved successfully",
                    "data": json.loads(dumps(list(minutes))),
                },
                200,
            )
        except Exception:
            return make_response({"success": "false", "message": "server error"}, 500)

    def generateminutesbyaudio(self, request):
        f = request.files["file"]
        path = "samples/audio/" + f.filename
        f.save(path)
        transcription = self.transcribe_audio(path)
        data = self.meeting_minutes(transcription)
        return make_response(
            {
                "success": "true",
                "message": "successfully generated minutes",
                "data": data,
            },
            200,
        )

    def generateminutesbytranscript(self, request):
        f = request.files["file"]
        path = "samples/transcripts/" + f.filename
        f.save(path)
        with open(path) as f:
            transcription = f.read()
        data = self.meeting_minutes(transcription)

        # data = transcription
        return make_response(
            {
                "success": "true",
                "message": "successfully generated minutes",
                "data": data,
            },
            200,
        )

    def saveminutes(self, request):
        # get the collections
        minutes = connect.minutes
        users = connect.users
        # check minutes format
        try:
            data = request.json
            Minutes(**data)
        except ValidationError as e:
            return make_response(
                {
                    "success": "false",
                    "message": ",".join([msg["msg"] for msg in e.errors()]),
                },
                200,
            )
        # add minutes with the user id
        # add the minute with the user id

        user = users.find_one({"email": request.user["email"]})
        data = request.json
        data["user_id"] = str(user["_id"])
        data["download_link"] = urllib.parse.quote(environ.get("BASE_URL")+"/docs/" + data["title"].split(":")[0]+".docx")
        minute_id = minutes.insert_one(data).inserted_id
        self.save_as_docx(request.json)
        # return the response
        return make_response(
            {"success": "true", "message": "successfully saved minutes","data":str(minute_id)}, 200
        )
    
    def deleteminute(self, id):
        try:
            minutes = connect.minutes
            minutes.delete_one({"_id":ObjectId(id)})
            return make_response({"success":"true","message":"document deleted successfully"},200)
        except:
            return make_response({"success":"false","message":"server error"},400)
        
    def updateminute(self, request):
        try:
            minutes = connect.minutes
            request["data"]["_id"] = ObjectId(request["data"]["_id"]["$oid"])
            minutes.find_one_and_update({"_id":ObjectId(request["id"])},{"$set":request["data"]})
            return make_response({"success":"true","message":"document updated successfully"},200)
        except:
            return make_response({"success":"false","message":"server error"},400)
